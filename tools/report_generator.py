#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
============================================================
报告生成器 V2 - 从测试数据 + 图片自动生成 Word 报告
============================================================
功能：
  - 从 CSV/JSON 测试数据文件读取测试结果
  - 自动插入测试图片（电路图、波形截图等）
  - 生成标准化 Word 文档（含表格、图表、封面）
  - 支持自定义报告模板和章节

依赖：pip install python-docx matplotlib pandas

用法：
  python report_generator_v2.py --data results.csv --images ./imgs/ -o report.docx
  python report_generator_v2.py --data results.json --template template.docx -o report.docx
============================================================
"""

import argparse
import csv
import json
import os
import sys
from datetime import datetime
from pathlib import Path

# ── 尝试导入可选依赖 ──────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    HAS_MPL = True
except ImportError:
    HAS_MPL = False


# ── 默认报告配置 ───────────────────────────────────────────
DEFAULT_CONFIG = {
    "title": "全国大学生电子设计竞赛 测试报告",
    "team": "参赛队",
    "problem": "赛题",
    "members": "队员1, 队员2, 队员3",
    "advisor": "指导教师",
    "school": "学校名称",
}


def load_test_data_csv(filepath: str) -> list[dict]:
    """
    从 CSV 文件加载测试数据。
    CSV 第一行为表头，后续行为数据。
    返回 list of dict。
    """
    data = []
    with open(filepath, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            data.append(dict(row))
    print(f"[INFO] 从 CSV 加载了 {len(data)} 条测试记录")
    return data


def load_test_data_json(filepath: str) -> list[dict]:
    """
    从 JSON 文件加载测试数据。
    支持格式：
      [{"name": ..., "value": ..., ...}, ...]
      或 {"tests": [...], "meta": {...}}
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        raw = json.load(f)

    if isinstance(raw, list):
        data = raw
    elif isinstance(raw, dict) and 'tests' in raw:
        data = raw['tests']
    else:
        data = [raw]

    print(f"[INFO] 从 JSON 加载了 {len(data)} 条测试记录")
    return data


def load_test_data(filepath: str) -> list[dict]:
    """根据文件扩展名自动选择加载方式。"""
    ext = Path(filepath).suffix.lower()
    if ext == '.csv':
        return load_test_data_csv(filepath)
    elif ext == '.json':
        return load_test_data_json(filepath)
    else:
        print(f"[WARN] 未知格式 {ext}，尝试按 CSV 解析")
        return load_test_data_csv(filepath)


def scan_images(image_dir: str) -> list[str]:
    """
    扫描目录中的图片文件，按文件名排序返回。
    支持 jpg, jpeg, png, bmp, gif, tiff。
    """
    if not os.path.isdir(image_dir):
        print(f"[WARN] 图片目录不存在: {image_dir}")
        return []

    exts = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'}
    images = []
    for f in sorted(os.listdir(image_dir)):
        if Path(f).suffix.lower() in exts:
            images.append(os.path.join(image_dir, f))

    print(f"[INFO] 扫描到 {len(images)} 张图片")
    return images


def add_cover_page(doc, config: dict):
    """添加封面页：标题、队伍信息、日期。"""
    # 空行留白
    for _ in range(4):
        doc.add_paragraph()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config.get('title', DEFAULT_CONFIG['title']))
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # 赛题信息
    info_lines = [
        f"赛题：{config.get('problem', DEFAULT_CONFIG['problem'])}",
        f"学校：{config.get('school', DEFAULT_CONFIG['school'])}",
        f"参赛队员：{config.get('members', DEFAULT_CONFIG['members'])}",
        f"指导教师：{config.get('advisor', DEFAULT_CONFIG['advisor'])}",
        f"日期：{datetime.now().strftime('%Y年%m月%d日')}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)

    # 分页
    doc.add_page_break()


def add_data_table(doc, data: list[dict], title: str = "测试数据"):
    """
    将测试数据以表格形式插入文档。
    自动从 dict 的 keys 中提取表头。
    """
    if not data:
        return

    doc.add_heading(title, level=2)

    headers = list(data[0].keys())
    table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            val = row_data.get(h, '')
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # 表后空行


def add_data_chart(doc, data: list[dict], chart_title: str = "测试数据趋势图"):
    """
    如果安装了 matplotlib，根据数据生成折线图并插入文档。
    自动选择第一列为 X 轴，数值列为 Y 轴。
    """
    if not HAS_MPL or not data:
        return

    headers = list(data[0].keys())
    if len(headers) < 2:
        return

    # 尝试找数值列
    x_col = headers[0]
    y_cols = []
    for h in headers[1:]:
        try:
            float(data[0][h])
            y_cols.append(h)
        except (ValueError, TypeError):
            pass

    if not y_cols:
        return

    fig, ax = plt.subplots(figsize=(8, 4))
    for y_col in y_cols:
        x_vals = list(range(len(data)))
        y_vals = []
        for row in data:
            try:
                y_vals.append(float(row[y_col]))
            except (ValueError, TypeError):
                y_vals.append(0)
        ax.plot(x_vals, y_vals, marker='o', label=y_col)

    ax.set_xlabel(x_col)
    ax.set_ylabel("数值")
    ax.set_title(chart_title)
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()

    # 保存临时图片
    tmp_path = os.path.join(os.path.dirname(doc.part.package.blob) if hasattr(doc, 'part') else '.', '_tmp_chart.png')
    tmp_path = '_tmp_chart.png'
    fig.savefig(tmp_path, dpi=150)
    plt.close(fig)

    doc.add_picture(tmp_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 清理临时文件
    try:
        os.remove(tmp_path)
    except OSError:
        pass


def add_images_section(doc, images: list[str]):
    """将图片逐一插入文档，每张图片下方标注文件名。"""
    if not images:
        return

    doc.add_heading("测试图片与波形", level=2)

    for img_path in images:
        fname = os.path.basename(img_path)
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 图片说明
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(f"图：{fname}")
            run.font.size = Pt(9)
            run.italic = True
            doc.add_paragraph()  # 空行
        except Exception as e:
            print(f"[WARN] 无法插入图片 {fname}: {e}")


def add_summary_section(doc, data: list[dict]):
    """根据测试数据自动生成简单统计摘要。"""
    doc.add_heading("测试结果统计摘要", level=2)

    if not data:
        doc.add_paragraph("无测试数据。")
        return

    headers = list(data[0].keys())
    summary_rows = []

    for h in headers:
        vals = []
        for row in data:
            try:
                vals.append(float(row[h]))
            except (ValueError, TypeError):
                continue
        if vals:
            summary_rows.append({
                "指标": h,
                "最小值": f"{min(vals):.4g}",
                "最大值": f"{max(vals):.4g}",
                "平均值": f"{sum(vals)/len(vals):.4g}",
                "数据量": str(len(vals)),
            })

    if summary_rows:
        add_data_table(doc, summary_rows, title="数值字段统计")


def generate_report(data: list[dict], images: list[str], output: str, config: dict):
    """主流程：组装并生成 Word 报告。"""
    if not HAS_DOCX:
        print("[ERROR] 缺少 python-docx 库，请运行: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)

    # ── 封面 ──
    add_cover_page(doc, config)

    # ── 目录占位 ──
    doc.add_heading("目录", level=1)
    doc.add_paragraph("（生成后请在 Word 中右键更新目录域）")
    doc.add_page_break()

    # ── 1. 系统概述 ──
    doc.add_heading("1. 系统概述", level=1)
    doc.add_paragraph(
        "本报告由报告生成器 V2 自动生成，包含系统测试数据、统计分析、"
        "以及相关测试图片。以下为自动生成的测试结果汇总。"
    )

    # ── 2. 测试数据 ──
    doc.add_heading("2. 测试数据", level=1)
    add_data_table(doc, data, title="测试记录表")
    add_data_chart(doc, data, chart_title="测试数据趋势图")

    # ── 3. 统计摘要 ──
    doc.add_heading("3. 测试结果统计", level=1)
    add_summary_section(doc, data)

    # ── 4. 测试图片 ──
    doc.add_heading("4. 测试图片与波形", level=1)
    add_images_section(doc, images)

    # ── 5. 结论 ──
    doc.add_heading("5. 结论", level=1)
    doc.add_paragraph("（请在此处补充测试结论与分析）")

    # ── 保存 ──
    doc.save(output)
    print(f"[OK] 报告已生成: {output}")
    print(f"     - 测试数据: {len(data)} 条")
    print(f"     - 图片: {len(images)} 张")


def main():
    parser = argparse.ArgumentParser(
        description='报告生成器 V2 - 从测试数据+图片自动生成Word报告',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python report_generator_v2.py --data results.csv --images ./imgs/ -o report.docx
  python report_generator_v2.py --data results.json -o report.docx
  python report_generator_v2.py --data results.csv --title "XX题测试报告" --team "A001"
        """
    )
    parser.add_argument('--data', '-d', required=True, help='测试数据文件 (CSV 或 JSON)')
    parser.add_argument('--images', '-i', default=None, help='图片目录路径')
    parser.add_argument('--output', '-o', default='report.docx', help='输出 Word 文件路径 (默认: report.docx)')
    parser.add_argument('--title', default=DEFAULT_CONFIG['title'], help='报告标题')
    parser.add_argument('--team', default=DEFAULT_CONFIG['team'], help='队伍编号')
    parser.add_argument('--problem', default=DEFAULT_CONFIG['problem'], help='赛题编号/名称')
    parser.add_argument('--members', default=DEFAULT_CONFIG['members'], help='参赛队员')
    parser.add_argument('--advisor', default=DEFAULT_CONFIG['advisor'], help='指导教师')
    parser.add_argument('--school', default=DEFAULT_CONFIG['school'], help='学校名称')

    args = parser.parse_args()

    # 检查依赖
    if not HAS_DOCX:
        print("[ERROR] 缺少 python-docx，请运行: pip install python-docx")
        sys.exit(1)

    # 加载测试数据
    if not os.path.isfile(args.data):
        print(f"[ERROR] 数据文件不存在: {args.data}")
        sys.exit(1)
    data = load_test_data(args.data)

    # 扫描图片
    images = scan_images(args.images) if args.images else []

    # 配置
    config = {
        "title": args.title,
        "team": args.team,
        "problem": args.problem,
        "members": args.members,
        "advisor": args.advisor,
        "school": args.school,
    }

    # 生成报告
    generate_report(data, images, args.output, config)


if __name__ == '__main__':
    main()
